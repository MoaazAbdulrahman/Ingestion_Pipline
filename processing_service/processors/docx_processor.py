from pathlib import Path
from typing import Dict, Any
from docx import Document

import sys
sys.path.append('/app/shared')
from logger import get_logger

logger = get_logger(__name__)


class DOCXProcessor:
    """
    DOCX document processor
    Extracts text content from DOCX files
    """
    
    def __init__(self):
        self.supported_extensions = ['.docx']
    
    def can_process(self, file_path: str) -> bool:
        """
        Check if file can be processed
        
        Args:
            file_path: Path to file
            
        Returns:
            True if file is DOCX
        """
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            logger.info(f"Processing DOCX: {file_path}")
            
            doc = Document(file_path)
            
            # Extract metadata from core properties
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "docx_metadata": {}
            }
            
            # Get document properties if available
            if doc.core_properties:
                props = doc.core_properties
                metadata["docx_metadata"] = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                }
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            
            if table_texts:
                full_text += "\n\n=== Tables ===\n\n" + "\n".join(table_texts)
            
            if not full_text.strip():
                raise ValueError("No text content extracted from DOCX")
            
            logger.info(
                f"Successfully extracted {len(full_text)} characters "
                f"from {len(paragraphs)} paragraphs and {len(doc.tables)} tables"
            )
            
            return {
                "text": full_text,
                "metadata": metadata,
                "paragraphs": paragraphs,
                "tables": table_texts
            }
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {str(e)}")
            raise
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validate DOCX file can be opened and read
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            True if file is valid
        """
        try:
            doc = Document(file_path)
            # Try to access paragraphs
            _ = doc.paragraphs
            return True
        except Exception as e:
            logger.error(f"Invalid DOCX file {file_path}: {str(e)}")
            return False